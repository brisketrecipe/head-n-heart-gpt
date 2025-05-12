import os
import io
from openai_service import OpenAIService
import docx

class DocumentProcessor:
    def __init__(self):
        self.openai_service = OpenAIService()

    def extract_content(self, file_path, file_content=None):
        """Extract text from various file types: PDF (PyPDF2, returns list of page texts), DOCX (python-docx), TXT (read), images (GPT-4o Vision)."""
        if file_content:
            file_content = io.BytesIO(file_content)
        file_extension = file_path.split('.')[-1].lower()
        if file_extension == 'pdf':
            import PyPDF2
            if file_content:
                file_content.seek(0)
                pdf_file = file_content
            else:
                pdf_file = open(file_path, 'rb')
            reader = PyPDF2.PdfReader(pdf_file)
            pages = [page.extract_text() or "" for page in reader.pages]
            if not file_content:
                pdf_file.close()
            return pages, 'text'
        elif file_extension in ['doc', 'docx']:
            import docx
            if file_content:
                file_content.seek(0)
                doc = docx.Document(file_content)
            else:
                doc = docx.Document(file_path)
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            return text, 'text'
        elif file_extension == 'txt':
            if file_content:
                file_content.seek(0)
                text = file_content.read().decode('utf-8')
            else:
                with open(file_path, 'r') as f:
                    text = f.read()
            return text, 'text'
        elif file_extension in ['jpg', 'jpeg', 'png']:
            if file_content:
                file_content.seek(0)
                text = self.openai_service.extract_text_from_image(file_content.read(), file_path)
                return text, 'text'
            else:
                with open(file_path, 'rb') as f:
                    text = self.openai_service.extract_text_from_image(f.read(), file_path)
                    return text, 'text'
        else:
            return f"Unsupported file type: {file_extension}", 'unsupported'

    def _extract_from_docx(self, source):
        """Extract text from Word documents"""
        doc = docx.Document(source)
        return "\n".join([paragraph.text for paragraph in doc.paragraphs])

    def _extract_from_txt(self, source):
        """Read plain text files"""
        if isinstance(source, str):
            with open(source, 'r') as file:
                return file.read()
        else:
            return source.read().decode('utf-8')

    def chunk_document(self, text, chunk_size=500, overlap=10):
        """Split document into chunks with overlap"""
        paragraphs = text.split('\n\n')
        chunks = []
        current_chunk = ""
        for paragraph in paragraphs:
            if not paragraph.strip():
                continue
            if len(current_chunk) + len(paragraph) > chunk_size and current_chunk:
                chunks.append(current_chunk.strip())
                current_chunk = current_chunk[-overlap:] if overlap > 0 else ""
            current_chunk += paragraph + "\n\n"
        if current_chunk.strip():
            chunks.append(current_chunk.strip())
        return chunks 